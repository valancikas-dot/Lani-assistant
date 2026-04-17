"""
Document reading, summarisation, and creation tools.

Supports:
  - Plain text (.txt, .md)
  - PDF (.pdf) via pypdf
  - Word documents (.docx) via python-docx
  - Excel spreadsheets (.xlsx) via openpyxl

Summarisation is a simple extractive placeholder that returns the first
N sentences.  Replace with an LLM call when an API key is available.

Creation tools:
  - CreateDocumentTool  → .docx (python-docx)
  - CreateSpreadsheetTool → .xlsx (openpyxl)
"""

import re
from pathlib import Path
from typing import Any, Dict, List

from app.schemas.commands import ToolResult
from app.tools.base import BaseTool
from app.tools.file_tools import _allowed, _deny

# ---------------------------------------------------------------------------
# Text extraction helpers
# ---------------------------------------------------------------------------

def _extract_text(path: str) -> str:
    """Extract plain text from txt, pdf, or docx files."""
    p = Path(path)
    suffix = p.suffix.lower()

    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="replace")

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError:
            raise RuntimeError("pypdf is not installed. Run: pip install pypdf")
        reader = PdfReader(str(p))
        return "\n".join(
            page.extract_text() or "" for page in reader.pages
        )

    if suffix == ".docx":
        try:
            from docx import Document  # type: ignore
        except ImportError:
            raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
        doc = Document(str(p))
        return "\n".join(para.text for para in doc.paragraphs)

    raise ValueError(f"Unsupported file type: {suffix}")


def _extractive_summary(text: str, max_sentences: int = 5) -> str:
    """
    Very simple extractive summary – returns the first *max_sentences* sentences.

    Replace this with an LLM-powered summary when an API key is available.
    """
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    summary_sentences = sentences[:max_sentences]
    return " ".join(summary_sentences)


# ---------------------------------------------------------------------------
# Tool implementations
# ---------------------------------------------------------------------------

class ReadDocumentTool(BaseTool):
    name = "read_document"
    description = "Read the full text content of a document (txt, pdf, docx)."
    requires_approval = False

    async def run(self, params: Dict[str, Any]) -> ToolResult:
        path: str = params["path"]
        if not _allowed(path):
            return _deny(path)
        if not Path(path).exists():
            return ToolResult(
                tool_name=self.name,
                status="error",
                message=f"File not found: {path}",
            )
        try:
            text = _extract_text(path)
            return ToolResult(
                tool_name=self.name,
                status="success",
                data={"path": path, "text": text, "char_count": len(text)},
                message=f"Read {len(text):,} characters from '{path}'.",
            )
        except Exception as exc:
            return ToolResult(tool_name=self.name, status="error", message=str(exc))


class SummarizeDocumentTool(BaseTool):
    name = "summarize_document"
    description = "Return a short summary of a document (txt, pdf, docx)."
    requires_approval = False

    async def run(self, params: Dict[str, Any]) -> ToolResult:
        path: str = params["path"]
        if not _allowed(path):
            return _deny(path)
        if not Path(path).exists():
            return ToolResult(
                tool_name=self.name,
                status="error",
                message=f"File not found: {path}",
            )
        try:
            text = _extract_text(path)
            summary = _extractive_summary(text)
            return ToolResult(
                tool_name=self.name,
                status="success",
                data={"path": path, "summary": summary},
                message="Summary generated (extractive – first 5 sentences).",
            )
        except Exception as exc:
            return ToolResult(tool_name=self.name, status="error", message=str(exc))


# ---------------------------------------------------------------------------
# Document creation tools
# ---------------------------------------------------------------------------

class CreateDocumentTool(BaseTool):
    """Create a Word document (.docx) from provided content.

    Params:
        path       – output file path ending in .docx (must be in allowed dirs)
        title      – document title (used as heading)
        content    – body text. Supports simple Markdown-like conventions:
                       Lines starting with '# ' become Heading1
                       Lines starting with '## ' become Heading2
                       Lines starting with '- ' become list items
                       Blank lines create paragraph breaks
        author     – optional author name for document properties
    """

    name = "create_document"
    description = (
        "Create a new Word (.docx) document with a title and body content. "
        "Content can contain Markdown-like headings and bullet lists."
    )
    requires_approval = False

    async def run(self, params: Dict[str, Any]) -> ToolResult:
        path: str = params.get("path", "")
        title: str = params.get("title", "Document")
        content: str = params.get("content", "")
        author: str = params.get("author", "Lani")

        if not path:
            return ToolResult(tool_name=self.name, status="error", message="'path' is required")
        if not path.lower().endswith(".docx"):
            return ToolResult(tool_name=self.name, status="error",
                              message="'path' must end with .docx")
        if not _allowed(path):
            return _deny(path)

        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt  # type: ignore
        except ImportError:
            return ToolResult(tool_name=self.name, status="error",
                              message="python-docx not installed. Run: pip install python-docx")

        try:
            p = Path(path)
            p.parent.mkdir(parents=True, exist_ok=True)

            doc = Document()
            # Document properties
            doc.core_properties.author = author
            doc.core_properties.title = title

            # Add title heading
            doc.add_heading(title, level=0)

            # Parse and render content lines
            for line in content.splitlines():
                stripped = line.strip()
                if not stripped:
                    doc.add_paragraph("")  # blank line spacing
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif re.match(r"^\d+\.\s", stripped):
                    doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
                else:
                    doc.add_paragraph(stripped)

            doc.save(str(p))
            size_kb = p.stat().st_size // 1024
            return ToolResult(
                tool_name=self.name,
                status="success",
                data={"path": str(p), "size_kb": size_kb},
                message=f"Dokumentas sukurtas: '{p.name}' ({size_kb} KB).",
            )
        except Exception as exc:
            return ToolResult(tool_name=self.name, status="error", message=str(exc))


class CreateSpreadsheetTool(BaseTool):
    """Create an Excel spreadsheet (.xlsx) from provided data.

    Params:
        path    – output file path ending in .xlsx (must be in allowed dirs)
        title   – sheet name (also used as a title row)
        headers – list of column header strings
        rows    – list of lists (each inner list is one data row)
        summary – optional text appended as a note below the data
    """

    name = "create_spreadsheet"
    description = (
        "Create a new Excel (.xlsx) spreadsheet with headers and rows of data. "
        "Provide 'headers' as a list of column names and 'rows' as a list of value lists."
    )
    requires_approval = False

    async def run(self, params: Dict[str, Any]) -> ToolResult:
        path: str = params.get("path", "")
        title: str = params.get("title", "Sheet1")
        headers: List[Any] = params.get("headers", [])
        rows: List[List[Any]] = params.get("rows", [])
        summary: str = params.get("summary", "")

        if not path:
            return ToolResult(tool_name=self.name, status="error", message="'path' is required")
        if not path.lower().endswith(".xlsx"):
            return ToolResult(tool_name=self.name, status="error",
                              message="'path' must end with .xlsx")
        if not _allowed(path):
            return _deny(path)
        if not headers:
            return ToolResult(tool_name=self.name, status="error",
                              message="'headers' list is required")

        try:
            import openpyxl  # type: ignore
            from openpyxl.styles import Font, PatternFill, Alignment  # type: ignore
        except ImportError:
            return ToolResult(tool_name=self.name, status="error",
                              message="openpyxl not installed. Run: pip install openpyxl")

        try:
            p = Path(path)
            p.parent.mkdir(parents=True, exist_ok=True)

            wb = openpyxl.Workbook()
            ws = wb.active
            assert ws is not None, "openpyxl Workbook.active is None"
            ws.title = title[:31]  # Excel sheet name limit

            # Header row styling
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill("solid", fgColor="4472C4")
            header_align = Alignment(horizontal="center", vertical="center")

            for col_idx, header in enumerate(headers, start=1):
                cell = ws.cell(row=1, column=col_idx, value=str(header))
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align

            # Data rows
            for row_idx, row_data in enumerate(rows, start=2):
                for col_idx, value in enumerate(row_data, start=1):
                    ws.cell(row=row_idx, column=col_idx, value=value)

            # Auto-size columns (approximate)
            for col in ws.columns:
                max_len = max(
                    (len(str(cell.value)) if cell.value is not None else 0 for cell in col),
                    default=10,
                )
                col_letter = col[0].column_letter if hasattr(col[0], "column_letter") else None  # type: ignore[union-attr]
                if col_letter:
                    ws.column_dimensions[col_letter].width = min(max_len + 4, 60)

            # Optional summary note below data
            if summary:
                note_row = len(rows) + 3
                ws.cell(row=note_row, column=1, value=summary).font = Font(italic=True)

            wb.save(str(p))
            size_kb = p.stat().st_size // 1024
            return ToolResult(
                tool_name=self.name,
                status="success",
                data={
                    "path": str(p),
                    "sheet": ws.title,
                    "rows": len(rows),
                    "columns": len(headers),
                    "size_kb": size_kb,
                },
                message=(
                    f"Lentelė sukurta: '{p.name}' "
                    f"({len(rows)} eilutės × {len(headers)} stulpeliai, {size_kb} KB)."
                ),
            )
        except Exception as exc:
            return ToolResult(tool_name=self.name, status="error", message=str(exc))
